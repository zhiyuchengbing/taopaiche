# -*- coding: utf-8 -*-
"""Generate 0522-0531 phase summary Word document."""
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from pathlib import Path


def set_run_font(run, name="宋体", size=12, bold=False, color=None):
    run.font.name = name
    run._element.rPr.rFonts.set(qn("w:eastAsia"), name)
    run.font.size = Pt(size)
    run.font.bold = bold
    if color:
        run.font.color.rgb = color


def add_heading(doc, text, level=1):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(12)
    p.paragraph_format.space_after = Pt(6)
    run = p.add_run(text)
    sizes = {1: 16, 2: 14, 3: 12}
    set_run_font(run, "黑体", sizes.get(level, 12), bold=True)


def add_body(doc, text, indent=False):
    p = doc.add_paragraph()
    p.paragraph_format.line_spacing = 1.5
    p.paragraph_format.space_after = Pt(6)
    if indent:
        p.paragraph_format.first_line_indent = Cm(0.74)
    run = p.add_run(text)
    set_run_font(run, "宋体", 12)
    return p


def add_bullet(doc, text, bold_prefix=""):
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.line_spacing = 1.5
    if bold_prefix:
        r1 = p.add_run(bold_prefix)
        set_run_font(r1, "宋体", 12, bold=True)
    r2 = p.add_run(text)
    set_run_font(r2, "宋体", 12)


def build_document():
    doc = Document()
    section = doc.sections[0]
    section.top_margin = Cm(2.54)
    section.bottom_margin = Cm(2.54)
    section.left_margin = Cm(3.17)
    section.right_margin = Cm(3.17)

    # Title
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    tr = title.add_run("0522—0531 阶段工作小结")
    set_run_font(tr, "黑体", 22, bold=True)

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sr = subtitle.add_run("Siamese + OCR + VLM 多级判别链路精度与可维护性收官优化")
    set_run_font(sr, "楷体", 14, color=RGBColor(64, 64, 64))

    meta = doc.add_paragraph()
    meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
    mr = meta.add_run("项目：过磅车辆套牌/换挂智能复核系统　　周期：2026年5月22日—5月31日")
    set_run_font(mr, "宋体", 10.5, color=RGBColor(96, 96, 96))

    doc.add_paragraph()

    add_heading(doc, "一、阶段定位与总体脉络", 1)
    add_body(
        doc,
        "本阶段工作承接 5 月上旬已建成的「Siamese 向量相似度 + 车头 OCR 预检 + 视觉语言模型（VLM）分部位二次复核」多级链路，"
        "由架构扩展全面转向判别精度治理与工程可维护性强化。在 OCR 抗干扰、触发门控松绑等中期改造已见成效的基础上，"
        "五月下旬（0522—0531）集中攻克多链路并行下的字段语义污染、提示词领域特征覆盖不足、模型输出解析脆弱性，"
        "以及尾部/头部多视角 AI 职责边界模糊等深层问题，完成了「字段解耦 → 提示词领域化 → 结构化输出 → 解析层护栏」"
        "的闭环迭代，将系统从「能跑通多链路」推进至「各链路可解释、可回退、可审计」的工程化水位。",
        indent=True,
    )

    add_heading(doc, "二、重点交付与核心突破", 1)

    add_heading(doc, "（一）多链路字段语义解耦与前端展示对齐", 2)
    add_bullet(doc, "彻底拆分主视角车尾 AI（tail1/tail2 裁切）与尾部视角 AI（3/4 原图）的结果字段，消除 `ai_tail_*` 与 `tail_second_check_*` 交叉写入导致的语义污染。", "【后端】")
    add_bullet(doc, "记录详情页 AI 字段命名与真实触发开关对齐（`main_tail_ai_used` 控制展示），差异卡片优先展示 `final_diff_summary`，避免 OCR 触发说明被误读为终裁结论。", "【前端】")
    add_bullet(doc, "头部视角裁切图改为 `contain` 完整展示，提升人工复核可用性与审计可追溯性。", "【体验】")

    add_heading(doc, "（二）提示词工程：分视角领域化与 Tier 结构化", 2)
    add_bullet(doc, "尾部视角车尾 AI 重构为 5 步 Tier-A/B 比对流水线（同位分区 → 后开口/侧围 → 颜色交叉校验 → 附属件 → 结论），号牌不可靠时强制走结构主链路，禁止空泛「栏板式一致」敷衍判读。", "【尾部视角】")
    add_bullet(doc, "建立「号牌/放大号清晰一致即正常」最高优先级短路规则，以及货物/篷布误判防护硬约束，显著降低货物标识、背景编号误触发换挂。", "【尾部视角】")
    add_bullet(doc, "主视角车尾提示词与尾部视角彻底分离，仅比对侧挡板、轴数、挡泥板、侧挂附件等本视角可见项，废弃不可见正后方部件的无效比对。", "【主视角车尾】")
    add_bullet(doc, "车头 AI 引入同位自证机制：文字/标识差异须声明子区域与双侧对齐状态，未对齐则文字证据无效；同步收紧导流罩/顶边灯/雨刷等易误判专节。", "【车头视角】")

    add_heading(doc, "（三）输出结构化与解析层护栏建设", 2)
    add_bullet(doc, "车头 AI 输出由「末行英文关键词」迁移为 JSON（`label` + `reason` 分离），根除 reason 正文中否定表述（如「不作为 fake_plate 依据」）被子串误匹配导致的反向误判。", "【0529—0531】")
    add_bullet(doc, "尾部视角 AI 新增成对可比对性审查（`pair_comparable`、`img*_trailer_rear_visible`），解析层 `_apply_comparability_rules` 对证据不足换挂自动降级为「无法判断」并回退主视角车尾 AI。", "【0531】")
    add_bullet(doc, "修复车头 AI 否定句式解析缺陷（「并非 fake_plate」类表述优先识别被否定项），补齐多关键词共存场景下的定案稳定性。", "【0525】")

    add_heading(doc, "（四）业务口径统一与可运维性提升", 2)
    add_bullet(doc, "移除套牌/换挂场景下 `最终差异总结` 的截断逻辑，终裁理由与 AI 判定理由保持一致，满足现场复核与留档审计需求。", "【0531】")
    add_bullet(doc, "清理 `predict_ai.py` 中已废弃的整车三分类、差异分析等冗余接口，降低代码熵与后续维护成本。", "【0529】")
    add_bullet(doc, "新增 `启动程序.bat` 一键启动脚本，缩短现场部署与调试路径。", "【0525】")

    add_heading(doc, "三、技术方法论总结", 1)
    add_body(
        doc,
        "本阶段形成可复用的「判别链路精细化治理」方法论，可概括为四条主线：",
        indent=True,
    )
    add_bullet(doc, "职责边界显性化——按视角（车头/主视角车尾/尾部视角）拆分提示词与结果字段，杜绝一条链路的结果污染另一条。")
    add_bullet(doc, "证据层级结构化——引入 Tier-A/B 与号牌短路规则，使 VLM 判读路径可复述、可质疑、可降级。")
    add_bullet(doc, "输出—解析解耦——JSON 结构化输出 + 后端 guard 规则，将「模型说了什么」与「系统采信什么」分离。")
    add_bullet(doc, "成像差异前置归因——成对可比对性、单侧阴影/眩光作废等规则，将工业现场光照不确定性纳入设计而非事后补丁。")

    add_heading(doc, "四、阶段成效（可量化维度）", 1)
    add_body(doc, "从工程交付视角，本阶段共涉及核心模块 4 个、提示词重构 3 套、解析/护栏逻辑 2 处、前端展示修正 3 项、冗余代码清理 4 个废弃接口，累计 README 变更条目 20+ 条。", indent=True)

    table = doc.add_table(rows=5, cols=2)
    table.style = "Table Grid"
    headers = [("优化维度", "预期效果"), ("字段解耦", "消除双尾部 AI 结果串写，记录页展示与真实触发一致"), ("提示词 Tier 化", "尾部换挂判读由「凭感觉」转为「按步骤、按层级」"), ("JSON + 解析护栏", "降低否定句式/子串误匹配类系统性误判"), ("成对可比对性", "3/4 视角拍不全挂车尾时自动回退，减少无效换挂告警")]
    for i, (a, b) in enumerate(headers):
        table.rows[i].cells[0].text = a
        table.rows[i].cells[1].text = b
        for cell in table.rows[i].cells:
            for para in cell.paragraphs:
                for run in para.runs:
                    set_run_font(run, "宋体", 10.5, bold=(i == 0))

    doc.add_paragraph()
    add_heading(doc, "五、阶段结语", 1)
    add_body(
        doc,
        "0522—0531 阶段标志着项目由「多级链路联通」迈入「多级链路可治理」新阶段。"
        "通过字段解耦、提示词领域化、结构化输出与解析护栏的系统性建设，"
        "VLM 在套牌/换挂判定中的角色进一步清晰：不再是黑盒终裁，而是可回退、可解释、受业务规则约束的复核节点。"
        "上述能力为 6 月继续深化的光照归因、几何优先比对、编号同类约束等精度优化奠定了工程底座。",
        indent=True,
    )

    out_dir = Path(r"d:\project\docs")
    out_dir.mkdir(parents=True, exist_ok=True)
    out = out_dir / "0522-0531_summary.docx"
    doc.save(str(out))
    print(out.resolve())


if __name__ == "__main__":
    build_document()
